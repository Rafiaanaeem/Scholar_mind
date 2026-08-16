import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

class ExportService:
    """Handles converting Markdown text into downloadable files (MD, DOCX, PDF)."""

    @staticmethod
    def export_to_markdown(content: str) -> bytes:
        """Converts text to bytes for a .md file download."""
        # We use .encode('utf-8') to turn the string into raw data bytes
        return content.encode('utf-8')

    @staticmethod
    def export_to_docx(topic: str, content: str) -> bytes:
        """Generates a Microsoft Word document in memory."""
        doc = Document()
        doc.add_heading(f"Research: {topic}", 0)
        
        # Basic parsing: split by newlines and add as paragraphs
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), level=2)
                else:
                    doc.add_paragraph(line)
                    
        # Save to an in-memory buffer instead of a real file on the hard drive
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0) # Reset the pointer to the beginning of the file
        return buffer.getvalue()

    @staticmethod
    def export_to_pdf(topic: str, content: str) -> bytes:
        """Generates a PDF document in memory using ReportLab."""
        buffer = io.BytesIO()
        # Create a document template
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []
        
        # Add Title
        title = Paragraph(f"<b>Research: {topic}</b>", styles['Title'])
        flowables.append(title)
        flowables.append(Spacer(1, 12))
        
        # Add Content
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('##'):
                    # Treat as header
                    p = Paragraph(f"<b>{line.replace('##', '').strip()}</b>", styles['Heading2'])
                else:
                    # Treat as normal text
                    p = Paragraph(line, styles['Normal'])
                flowables.append(p)
                flowables.append(Spacer(1, 6))
                
        doc.build(flowables)
        buffer.seek(0)
        return buffer.getvalue()